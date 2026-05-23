"""Regenerate Section 4.2 tables, figures, and Word content.

This script is intentionally self-contained because the project environment and
the document-editing environment may not expose the same Python packages.

Typical usage:

    conda run -n tsf-imputation-analysis python tools/rework_section_4_2_results.py --figures-only
    python tools/rework_section_4_2_results.py --docx-only

The first command uses the project environment to generate data summaries and
figures. The second command uses python-docx, if available, to replace the
current Section 4.2 in the manuscript.
"""

from __future__ import annotations

import argparse
import csv
import math
import statistics
from collections import Counter, defaultdict
from pathlib import Path
from typing import Dict, Iterable, List, Sequence


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = REPO_ROOT / "results_analysis" / "块状缺失对长序列预测影响-实验结果统计-0509.csv"
DEFAULT_DOCX = REPO_ROOT / "文档资料" / "面向时序基础模型预测的缺失填补影响分析-muyun-260523.docx"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "figures" / "section_4_2"
DEFAULT_DATA_DIR = REPO_ROOT / "figures" / "data"
DATA_MANIFEST = REPO_ROOT / "figures" / "data-manifest.md"

METHOD_LABELS = {
    "/": "Clean",
    "均值": "Mean",
    "前项": "Forward",
    "后向": "Backward",
    "线性": "Linear",
    "kalman_struct": "Kalman-Struct",
    "kalman_arima": "Kalman-ARIMA",
    "gp_rbf": "GP-RBF",
    "saits": "SAITS",
}

METHOD_ORDER = [
    "SAITS",
    "GP-RBF",
    "Kalman-ARIMA",
    "Linear",
    "Kalman-Struct",
    "Backward",
    "Forward",
    "Mean",
]

MODEL_LABELS = {
    "visiontspp": "VisionTS++",
    "kairos50m": "Kairos-50M",
    "kairos23m": "Kairos-23M",
    "timesfm2p5": "TimesFM-2.5",
    "sundial": "Sundial",
    "chronos2": "Chronos-2",
    "timesfm2p0": "TimesFM-2.0",
}

HISTORY_METRICS = [
    "trend_strength",
    "trend_linearity",
    "seasonal_strength",
    "seasonal_correlation",
    "residual_autocorr",
    "spectral_entropy",
]

HISTORY_METRIC_LABELS = [
    "趋势强度",
    "趋势线性度",
    "季节强度",
    "季节相关性",
    "残差自相关性",
    "谱熵",
]

MODEL_ORDER = [
    "VisionTS++",
    "Kairos-50M",
    "Kairos-23M",
    "TimesFM-2.5",
    "Sundial",
    "Chronos-2",
    "TimesFM-2.0",
]


def parse_float(value: object) -> float | None:
    text = "" if value is None else str(value).strip()
    if not text or text == "/":
        return None
    try:
        return float(text)
    except ValueError:
        return None


def median(values: Iterable[float | None]) -> float:
    clean = [v for v in values if v is not None and not math.isnan(v)]
    return statistics.median(clean)


def quantile(values: Iterable[float | None], q: float) -> float:
    clean = sorted(v for v in values if v is not None and not math.isnan(v))
    if not clean:
        return float("nan")
    pos = (len(clean) - 1) * q
    lo = math.floor(pos)
    hi = math.ceil(pos)
    if lo == hi:
        return clean[lo]
    return clean[lo] * (hi - pos) + clean[hi] * (pos - lo)


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{value * 100:.{digits}f}%"


def fmt_num(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}"


def load_results(path: Path) -> List[dict]:
    """Read the two-row-header CSV and forward-fill model, dataset, and ratio.

    The source table stores repeated model/dataset/ratio values as blanks. This
    parser mirrors the visual table layout and creates one normalized record per
    row.
    """

    rows: List[dict] = []
    current = {"model": "", "dataset": "", "ratio": ""}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        next(reader)
        next(reader)
        for raw in reader:
            if len(raw) < 19 or not any(raw):
                continue
            if raw[0].strip():
                current["model"] = raw[0].strip()
            if raw[1].strip():
                current["dataset"] = raw[1].strip()
            if raw[2].strip():
                current["ratio"] = raw[2].strip()
            method = raw[3].strip()
            if not method:
                continue

            record = {
                "model_raw": current["model"],
                "model": MODEL_LABELS.get(current["model"], current["model"]),
                "dataset": current["dataset"],
                "missing_ratio": parse_float(current["ratio"]),
                "method_raw": method,
                "method": METHOD_LABELS.get(method, method),
                "imputation_smape": parse_float(raw[4]),
                "forecast_mse": parse_float(raw[5]),
                "forecast_smape": parse_float(raw[6]),
            }
            for idx, metric in enumerate(HISTORY_METRICS):
                record[f"history_{metric}"] = parse_float(raw[7 + idx])
            for idx, metric in enumerate(HISTORY_METRICS):
                record[f"prediction_{metric}"] = parse_float(raw[13 + idx])
            rows.append(record)
    return rows


def attach_baseline(rows: Sequence[dict]) -> List[dict]:
    clean = {
        (row["model"], row["dataset"]): row
        for row in rows
        if row["missing_ratio"] == 0 and row["method"] == "Clean"
    }

    records: List[dict] = []
    for row in rows:
        ratio = row["missing_ratio"]
        if ratio is None or ratio <= 0:
            continue
        baseline = clean.get((row["model"], row["dataset"]))
        if baseline is None:
            continue
        if not baseline["forecast_mse"] or not baseline["forecast_smape"]:
            continue

        record = dict(row)
        record["clean_forecast_mse"] = baseline["forecast_mse"]
        record["clean_forecast_smape"] = baseline["forecast_smape"]
        record["relative_mse_gain"] = (
            row["forecast_mse"] - baseline["forecast_mse"]
        ) / abs(baseline["forecast_mse"])
        record["relative_smape_gain"] = (
            row["forecast_smape"] - baseline["forecast_smape"]
        ) / abs(baseline["forecast_smape"])
        record["mse_better_than_clean"] = row["forecast_mse"] < baseline["forecast_mse"]

        for metric in HISTORY_METRICS:
            value = row[f"history_{metric}"]
            base = baseline[f"history_{metric}"]
            record[f"delta_{metric}"] = abs(value - base) if value is not None and base is not None else None
        records.append(record)

    add_structure_drift(records)
    return records


def add_structure_drift(records: List[dict]) -> None:
    """Add a scale-balanced history structure drift score.

    Each history metric has a different numerical scale. The script divides the
    absolute drift of each channel by its 95th percentile across all imputed
    records, then averages the six normalized channels. This keeps spectral
    entropy from dominating the combined score.
    """

    denominators: Dict[str, float] = {}
    for metric in HISTORY_METRICS:
        values = sorted(
            row[f"delta_{metric}"]
            for row in records
            if row[f"delta_{metric}"] is not None
        )
        idx = int(0.95 * (len(values) - 1))
        denominators[metric] = values[idx] if values and values[idx] else 1.0

    for row in records:
        parts = []
        for metric in HISTORY_METRICS:
            delta = row[f"delta_{metric}"]
            if delta is not None:
                parts.append(delta / denominators[metric])
        row["history_structure_drift"] = sum(parts) / len(parts)


def group_by(records: Sequence[dict], key: str) -> Dict[str, List[dict]]:
    grouped: Dict[str, List[dict]] = defaultdict(list)
    for record in records:
        grouped[record[key]].append(record)
    return grouped


def summarize_model(records: Sequence[dict]) -> List[dict]:
    summary = []
    for model, group in group_by(records, "model").items():
        summary.append(
            {
                "基础模型": model,
                "记录数": len(group),
                "MSE相对增幅中位数": fmt_pct(median(r["relative_mse_gain"] for r in group)),
                "MSE相对增幅90分位数": fmt_pct(quantile((r["relative_mse_gain"] for r in group), 0.90)),
                "sMAPE相对增幅中位数": fmt_pct(median(r["relative_smape_gain"] for r in group)),
                "MSE低于clean比例": fmt_pct(sum(r["mse_better_than_clean"] for r in group) / len(group)),
            }
        )
    return sorted(summary, key=lambda r: MODEL_ORDER.index(r["基础模型"]))


def summarize_method(records: Sequence[dict]) -> List[dict]:
    best_counts = best_method_counts(records)
    summary = []
    for method in METHOD_ORDER:
        group = [r for r in records if r["method"] == method]
        summary.append(
            {
                "填补方法": method,
                "填补SMAPE中位数": fmt_num(median(r["imputation_smape"] for r in group), 2),
                "历史结构扰动中位数": fmt_num(median(r["history_structure_drift"] for r in group), 3),
                "MSE相对增幅中位数": fmt_pct(median(r["relative_mse_gain"] for r in group)),
                "sMAPE相对增幅中位数": fmt_pct(median(r["relative_smape_gain"] for r in group)),
                "预测最优次数": best_counts.get(method, 0),
            }
        )
    return summary


def summarize_ratio(records: Sequence[dict]) -> List[dict]:
    summary = []
    for ratio in [0.1, 0.2, 0.3]:
        group = [r for r in records if abs(r["missing_ratio"] - ratio) < 1e-9]
        summary.append(
            {
                "缺失率": f"{int(ratio * 100)}%",
                "填补SMAPE中位数": fmt_num(median(r["imputation_smape"] for r in group), 2),
                "历史结构扰动中位数": fmt_num(median(r["history_structure_drift"] for r in group), 3),
                "MSE相对增幅中位数": fmt_pct(median(r["relative_mse_gain"] for r in group)),
                "sMAPE相对增幅中位数": fmt_pct(median(r["relative_smape_gain"] for r in group)),
            }
        )
    return summary


def summarize_dataset(records: Sequence[dict]) -> List[dict]:
    grouped = group_by(records, "dataset")
    rows = []
    for dataset, group in grouped.items():
        rows.append(
            {
                "数据集": dataset,
                "MSE相对增幅中位数": fmt_pct(median(r["relative_mse_gain"] for r in group)),
                "sMAPE相对增幅中位数": fmt_pct(median(r["relative_smape_gain"] for r in group)),
                "历史结构扰动中位数": fmt_num(median(r["history_structure_drift"] for r in group), 3),
                "_sort": median(r["relative_mse_gain"] for r in group),
            }
        )
    rows = sorted(rows, key=lambda r: r["_sort"], reverse=True)
    selected = rows[:6] + rows[-5:]
    for row in selected:
        row.pop("_sort", None)
    return selected


def best_method_counts(records: Sequence[dict]) -> Counter:
    grouped: Dict[tuple, List[dict]] = defaultdict(list)
    for record in records:
        grouped[(record["model"], record["dataset"], record["missing_ratio"])].append(record)

    counts: Counter = Counter()
    for group in grouped.values():
        counts[min(group, key=lambda r: r["relative_mse_gain"])["method"]] += 1
    return counts


def monotonic_ratio_summary(records: Sequence[dict]) -> dict:
    grouped: Dict[tuple, Dict[float, float]] = defaultdict(dict)
    for record in records:
        grouped[(record["model"], record["dataset"], record["method"])][record["missing_ratio"]] = record[
            "relative_mse_gain"
        ]

    increasing = decreasing = non_monotonic = 0
    for values in grouped.values():
        if not all(ratio in values for ratio in [0.1, 0.2, 0.3]):
            continue
        a, b, c = values[0.1], values[0.2], values[0.3]
        if a < b < c:
            increasing += 1
        elif a > b > c:
            decreasing += 1
        else:
            non_monotonic += 1
    total = increasing + decreasing + non_monotonic
    return {
        "strict_increasing": increasing,
        "strict_decreasing": decreasing,
        "non_monotonic": non_monotonic,
        "total": total,
        "strict_increasing_ratio": increasing / total,
        "strict_decreasing_ratio": decreasing / total,
        "non_monotonic_ratio": non_monotonic / total,
    }


def rankdata(values: Sequence[float]) -> List[float]:
    pairs = sorted((value, idx) for idx, value in enumerate(values))
    ranks = [0.0] * len(values)
    pos = 0
    while pos < len(pairs):
        end = pos
        while end + 1 < len(pairs) and pairs[end + 1][0] == pairs[pos][0]:
            end += 1
        avg_rank = (pos + end + 2) / 2.0
        for item in range(pos, end + 1):
            ranks[pairs[item][1]] = avg_rank
        pos = end + 1
    return ranks


def pearson(x_values: Sequence[float], y_values: Sequence[float]) -> float | None:
    if len(x_values) != len(y_values) or not x_values:
        return None
    x_mean = sum(x_values) / len(x_values)
    y_mean = sum(y_values) / len(y_values)
    x_delta = [x - x_mean for x in x_values]
    y_delta = [y - y_mean for y in y_values]
    x_norm = math.sqrt(sum(x * x for x in x_delta))
    y_norm = math.sqrt(sum(y * y for y in y_delta))
    if not x_norm or not y_norm:
        return None
    return sum(x * y for x, y in zip(x_delta, y_delta)) / (x_norm * y_norm)


def spearman(x_values: Sequence[float], y_values: Sequence[float]) -> float | None:
    return pearson(rankdata(x_values), rankdata(y_values))


def correlation_summary(records: Sequence[dict]) -> dict:
    grouped: Dict[tuple, List[dict]] = defaultdict(list)
    for record in records:
        grouped[(record["model"], record["dataset"], record["missing_ratio"])].append(record)

    out = {}
    for x_col in ["imputation_smape", "history_structure_drift"]:
        correlations = []
        for group in grouped.values():
            x_values = [r[x_col] for r in group]
            y_values = [r["relative_mse_gain"] for r in group]
            if len(set(x_values)) <= 1 or len(set(y_values)) <= 1:
                continue
            rho = spearman(x_values, y_values)
            if rho is not None:
                correlations.append(rho)
        out[x_col] = {
            "group_count": len(correlations),
            "median_spearman": median(correlations),
            "positive_ratio": sum(value > 0 for value in correlations) / len(correlations),
            "overall_spearman": spearman(
                [r[x_col] for r in records],
                [r["relative_mse_gain"] for r in records],
            ),
        }
    return out


def write_csv(path: Path, rows: Sequence[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        return
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_data_outputs(records: Sequence[dict], output_dir: Path, data_dir: Path) -> dict:
    data_dir.mkdir(parents=True, exist_ok=True)
    outputs = {
        "model_summary": data_dir / "section_4_2_model_summary.csv",
        "method_summary": data_dir / "section_4_2_method_summary.csv",
        "ratio_summary": data_dir / "section_4_2_missing_ratio_summary.csv",
        "dataset_summary": data_dir / "section_4_2_dataset_sensitivity_summary.csv",
    }
    write_csv(outputs["model_summary"], summarize_model(records))
    write_csv(outputs["method_summary"], summarize_method(records))
    write_csv(outputs["ratio_summary"], summarize_ratio(records))
    write_csv(outputs["dataset_summary"], summarize_dataset(records))
    return outputs


def setup_matplotlib():
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager

    # Windows Chinese fonts are tried first; matplotlib will fall back if absent.
    candidates = ["Microsoft YaHei", "SimHei", "SimSun", "Arial Unicode MS", "DejaVu Sans"]
    installed = {font.name for font in font_manager.fontManager.ttflist}
    for name in candidates:
        if name in installed:
            plt.rcParams["font.sans-serif"] = [name]
            break
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams.update(
        {
            "font.size": 9,
            "axes.titlesize": 11,
            "axes.labelsize": 9,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.grid": True,
            "grid.alpha": 0.25,
            "legend.frameon": False,
            "savefig.dpi": 450,
            "savefig.bbox": "tight",
        }
    )
    return plt


def plot_heatmap(records: Sequence[dict], output_dir: Path) -> Path:
    plt = setup_matplotlib()
    import numpy as np
    from matplotlib.colors import TwoSlopeNorm

    matrix = []
    for model in MODEL_ORDER:
        row = []
        for method in METHOD_ORDER:
            group = [r for r in records if r["model"] == model and r["method"] == method]
            row.append(median(r["relative_mse_gain"] for r in group) * 100)
        matrix.append(row)
    values = np.array(matrix)

    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    limit = max(abs(float(np.nanmin(values))), abs(float(np.nanmax(values))))
    norm = TwoSlopeNorm(vmin=-limit, vcenter=0, vmax=limit)
    image = ax.imshow(values, cmap="RdBu_r", norm=norm, aspect="auto")
    ax.set_xticks(range(len(METHOD_ORDER)))
    ax.set_xticklabels(METHOD_ORDER, rotation=35, ha="right")
    ax.set_yticks(range(len(MODEL_ORDER)))
    ax.set_yticklabels(MODEL_ORDER)
    ax.set_title("不同模型与填补方法组合下的MSE相对增幅中位数")
    for i in range(values.shape[0]):
        for j in range(values.shape[1]):
            color = "white" if abs(values[i, j]) > limit * 0.55 else "black"
            ax.text(j, i, f"{values[i, j]:.1f}%", ha="center", va="center", color=color, fontsize=7)
    colorbar = fig.colorbar(image, ax=ax, fraction=0.035, pad=0.02)
    colorbar.set_label("MSE相对增幅中位数")
    output = output_dir / "fig4_2_1_model_imputer_heatmap.png"
    fig.savefig(output)
    fig.savefig(output.with_suffix(".svg"))
    plt.close(fig)
    return output


def plot_missing_ratio(records: Sequence[dict], output_dir: Path) -> Path:
    plt = setup_matplotlib()
    ratio_labels = ["10%", "20%", "30%"]
    ratio_values = [0.1, 0.2, 0.3]
    data = [
        [r["relative_mse_gain"] * 100 for r in records if abs(r["missing_ratio"] - ratio) < 1e-9]
        for ratio in ratio_values
    ]
    medians = [statistics.median(values) for values in data]

    fig, ax = plt.subplots(figsize=(6.5, 4.2))
    box = ax.boxplot(data, tick_labels=ratio_labels, showfliers=False, patch_artist=True)
    colors = ["#76B7B2", "#F28E2B", "#E15759"]
    for patch, color in zip(box["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_alpha(0.75)
        patch.set_edgecolor("#333333")
    for median_line in box["medians"]:
        median_line.set_color("#111111")
        median_line.set_linewidth(1.6)
    ax.plot([1, 2, 3], medians, color="#4E79A7", marker="o", linewidth=1.8, label="中位数")
    for idx, value in enumerate(medians, start=1):
        ax.text(idx, value + 2.0, f"{value:.1f}%", ha="center", va="bottom", fontsize=8)
    ax.axhline(0, color="#555555", linewidth=0.8)
    ax.set_xlabel("缺失率")
    ax.set_ylabel("MSE相对增幅")
    ax.set_title("缺失率提升下的预测误差变化分布")
    ax.legend(loc="upper left")
    output = output_dir / "fig4_2_2_missing_ratio_distribution.png"
    fig.savefig(output)
    fig.savefig(output.with_suffix(".svg"))
    plt.close(fig)
    return output


def plot_structure_relation(records: Sequence[dict], output_dir: Path, corr: dict) -> Path:
    plt = setup_matplotlib()

    fig, axes = plt.subplots(1, 2, figsize=(9.0, 3.9), sharey=True)
    specs = [
        (
            "imputation_smape",
            "填补SMAPE",
            "#4E79A7",
            corr["imputation_smape"],
        ),
        (
            "history_structure_drift",
            "历史结构扰动",
            "#E15759",
            corr["history_structure_drift"],
        ),
    ]
    y_values = [r["relative_mse_gain"] * 100 for r in records]
    y_lower = quantile(y_values, 0.005) - 10
    y_upper = quantile(y_values, 0.99) + 20
    for ax, (x_col, x_label, color, stats) in zip(axes, specs):
        x_values = [r[x_col] for r in records]
        ax.scatter(x_values, y_values, s=9, alpha=0.18, color=color, edgecolors="none")
        ax.axhline(0, color="#555555", linewidth=0.8)
        ax.set_ylim(y_lower, y_upper)
        ax.set_xlabel(x_label)
        ax.text(
            0.03,
            0.95,
            f"分组Spearman中位数={stats['median_spearman']:.3f}\n正相关组占比={stats['positive_ratio'] * 100:.1f}%",
            transform=ax.transAxes,
            va="top",
            fontsize=8,
            bbox={"boxstyle": "round,pad=0.25", "facecolor": "white", "edgecolor": "#cccccc", "alpha": 0.9},
        )
    axes[0].set_ylabel("MSE相对增幅")
    axes[0].set_title("点级重构误差与下游误差")
    axes[1].set_title("结构扰动与下游误差")
    fig.subplots_adjust(bottom=0.20, wspace=0.20)
    fig.text(0.5, 0.04, "注：纵轴显示至约99分位，以避免极端点压缩主体分布。", ha="center", fontsize=8)
    output = output_dir / "fig4_2_3_error_relation.png"
    fig.savefig(output)
    fig.savefig(output.with_suffix(".svg"))
    plt.close(fig)
    return output


def write_manifest(data_outputs: dict, figure_outputs: Sequence[Path], input_path: Path) -> None:
    DATA_MANIFEST.parent.mkdir(parents=True, exist_ok=True)
    rows = [
        "| Figure/Table | Data file | Real/mock | Source | Script | Outputs |",
        "|---|---|---|---|---|---|",
    ]
    for name, data_path in data_outputs.items():
        rows.append(
            f"| {name} | {data_path.relative_to(REPO_ROOT)} | real | {input_path.relative_to(REPO_ROOT)} | tools/rework_section_4_2_results.py | {data_path.relative_to(REPO_ROOT)} |"
        )
    for figure in figure_outputs:
        rows.append(
            f"| {figure.stem} | figures/data/section_4_2_*.csv | real | {input_path.relative_to(REPO_ROOT)} | tools/rework_section_4_2_results.py | {figure.relative_to(REPO_ROOT)}; {figure.with_suffix('.svg').relative_to(REPO_ROOT)} |"
        )
    DATA_MANIFEST.write_text("\n".join(rows) + "\n", encoding="utf-8")


def generate_outputs(input_path: Path, output_dir: Path, data_dir: Path) -> dict:
    rows = load_results(input_path)
    records = attach_baseline(rows)
    output_dir.mkdir(parents=True, exist_ok=True)
    data_outputs = write_data_outputs(records, output_dir, data_dir)
    corr = correlation_summary(records)
    figure_outputs = [
        plot_heatmap(records, output_dir),
        plot_missing_ratio(records, output_dir),
        plot_structure_relation(records, output_dir, corr),
    ]
    write_manifest(data_outputs, figure_outputs, input_path)
    return {
        "rows": rows,
        "records": records,
        "data_outputs": data_outputs,
        "figure_outputs": figure_outputs,
        "correlation": corr,
        "monotonic": monotonic_ratio_summary(records),
        "best_counts": best_method_counts(records),
    }


def table_from_rows(doc, rows: Sequence[dict], columns: Sequence[str]):
    table = doc.add_table(rows=1, cols=len(columns))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header = table.rows[0].cells
    for idx, col in enumerate(columns):
        header[idx].text = col
    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            cells[idx].text = str(row[col])
    return table


def move_block_to(body, element, index: int) -> int:
    element.getparent().remove(element)
    body.insert(index, element)
    return index + 1


def add_moved_paragraph(doc, body, index: int, text: str, style=None):
    para = doc.add_paragraph(text, style=style)
    return move_block_to(body, para._element, index)


def add_moved_table(doc, body, index: int, rows: Sequence[dict], columns: Sequence[str]):
    table = table_from_rows(doc, rows, columns)
    return move_block_to(body, table._element, index)


def add_moved_picture(doc, body, index: int, path: Path, width_inches: float = 6.4):
    from docx.shared import Inches

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    return move_block_to(body, para._element, index)


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))


def find_body_index(body, prefix: str) -> int:
    for idx, element in enumerate(body):
        if element_text(element).strip().startswith(prefix):
            return idx
    raise ValueError(f"Cannot find body element starting with: {prefix}")


def update_docx(docx_path: Path, output_dir: Path, input_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required for --docx-only. Run this part with an interpreter that has python-docx installed."
        ) from exc

    rows = load_results(input_path)
    records = attach_baseline(rows)
    model_rows = summarize_model(records)
    method_rows = summarize_method(records)
    ratio_rows = summarize_ratio(records)
    dataset_rows = summarize_dataset(records)
    corr = correlation_summary(records)
    monotonic = monotonic_ratio_summary(records)
    best_counts = best_method_counts(records)

    doc = Document(str(docx_path))
    body = doc._body._element
    start = find_body_index(body, "4.2 总体实验结果分析")
    end = find_body_index(body, "4.3 ")

    heading_style = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("4.2 总体实验结果分析"):
            heading_style = paragraph.style
            break

    for idx in range(end - 1, start - 1, -1):
        body.remove(body[idx])

    insert_at = start
    insert_at = add_moved_paragraph(doc, body, insert_at, "4.2 总体实验结果分析", style=heading_style)
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "本节基于主实验汇总表进行聚合分析。原始记录以基础模型、数据集、缺失率和填补方法为单位，"
            f"包含{len(set(r['model'] for r in records))}个基础模型、{len(set(r['dataset'] for r in records))}个数据集、"
            "3个缺失率和8种填补方法，共4032条填补预测记录，并以168条clean输入记录作为同模型同数据集下的参照。"
            "下文以相对MSE增幅作为主指标，以相对sMAPE增幅作为辅助指标。历史结构扰动由六个回顾窗口结构通道构成，"
            "每个通道先计算相对clean输入的绝对漂移，再按全体记录中该通道漂移的95分位数归一化，最后取六个通道的平均值。"
        ),
    )

    insert_at = add_moved_paragraph(doc, body, insert_at, "表4.2.1 基础模型鲁棒性汇总")
    insert_at = add_moved_table(
        doc,
        body,
        insert_at,
        model_rows,
        ["基础模型", "MSE相对增幅中位数", "MSE相对增幅90分位数", "sMAPE相对增幅中位数", "MSE低于clean比例"],
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "表4.2.1显示，块状缺失经填补后对不同基础模型的影响差异较大。TimesFM-2.0的MSE相对增幅中位数为14.0%，"
            "90分位数达到122.7%，说明其在当前设置下更容易出现尾部误差放大。Chronos-2、Sundial和TimesFM-2.5的中位增幅"
            "位于6.0%到8.0%之间，Kairos-23M与Kairos-50M分别为4.9%和3.7%。VisionTS++是主要例外，其MSE相对变化中位数为"
            "-3.2%，且65.8%的填补记录低于clean输入误差，说明平滑或结构压缩型输入在该模型上可能与其预测偏好相匹配。"
        ),
    )

    insert_at = add_moved_picture(doc, body, insert_at, output_dir / "fig4_2_1_model_imputer_heatmap.png")
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        "图4.2.1 不同基础模型与填补方法组合下的MSE相对增幅中位数",
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "图4.2.1进一步表明，填补方法的优劣依赖于具体基础模型。按整体中位数排序，SAITS和GP-RBF的MSE相对增幅分别为"
            "2.5%和2.6%，是总体表现较稳定的两类方法；Linear的填补SMAPE中位数最低，为2.12，但其预测误差增幅中位数为4.0%，"
            "说明点级重构精度并不能直接等同于下游预测收益。Mean的整体误差增幅最高，为6.8%，但它仍有"
            f"{best_counts.get('Mean', 0)}次成为同一模型、数据集和缺失率组合下的预测最优方法，其中VisionTS++贡献较多。"
            "这一结果支持模型感知的填补选择，即填补算法评价需要纳入下游基础模型。"
        ),
    )

    insert_at = add_moved_picture(doc, body, insert_at, output_dir / "fig4_2_2_missing_ratio_distribution.png", width_inches=5.6)
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        "图4.2.2 不同缺失率下的MSE相对增幅分布",
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "缺失率提高会整体放大预测误差。10%、20%和30%缺失率下的MSE相对增幅中位数分别为2.2%、4.8%和8.8%，"
            "填补SMAPE中位数也由1.35升至4.71，历史结构扰动中位数由0.090升至0.305。局部组合中仍存在非单调现象，"
            f"在同一模型、数据集和填补方法内，仅{monotonic['strict_increasing_ratio'] * 100:.1f}%的序列表现为10%、20%、30%严格递增，"
            f"{monotonic['non_monotonic_ratio'] * 100:.1f}%为非单调，{monotonic['strict_decreasing_ratio'] * 100:.1f}%为严格递减。"
            "因此，缺失率更适合作为总体风险放大因素，而具体误差变化还受缺失位置、局部序列形态、填补平滑效应和模型外推偏好共同影响。"
        ),
    )

    insert_at = add_moved_paragraph(doc, body, insert_at, "表4.2.2 数据集敏感性汇总")
    insert_at = add_moved_table(
        doc,
        body,
        insert_at,
        dataset_rows,
        ["数据集", "MSE相对增幅中位数", "sMAPE相对增幅中位数", "历史结构扰动中位数"],
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "表4.2.2列出了MSE相对增幅最高的六个数据集和最低的五个数据集。traffic和Finland_Traffic_15T的中位增幅分别为"
            "70.4%和59.3%，明显高于其他数据集，说明强同步、强周期且clean误差较低的交通类序列更容易受到历史窗口结构扰动影响。"
            "azure2019_U_5T、Port_Activity_D、Coastal_T_S_15T、current_velocity_H和weather的中位变化接近0或低于0，"
            "说明部分数据条件下填补引入的平滑变化可以被模型吸收。"
        ),
    )

    insert_at = add_moved_picture(doc, body, insert_at, output_dir / "fig4_2_3_error_relation.png")
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        "图4.2.3 点级填补误差、历史结构扰动与下游预测误差的关系（纵轴显示至约99分位）",
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "图4.2.3比较了点级填补SMAPE和历史结构扰动对下游误差变化的解释能力。在固定模型、数据集和缺失率后，"
            f"填补SMAPE与MSE相对增幅的分组Spearman相关系数中位数为{corr['imputation_smape']['median_spearman']:.3f}，"
            f"正相关组占{corr['imputation_smape']['positive_ratio'] * 100:.1f}%；历史结构扰动的对应中位数为"
            f"{corr['history_structure_drift']['median_spearman']:.3f}，正相关组占{corr['history_structure_drift']['positive_ratio'] * 100:.1f}%。"
            "这表明点级重构误差包含有效信息，历史窗口结构变化与下游预测误差之间的关联更稳定。"
        ),
    )
    insert_at = add_moved_paragraph(
        doc,
        body,
        insert_at,
        (
            "总体来看，主实验结果可归纳为三个方面。第一，块状缺失经填补后通常会提高预测误差，但影响强度首先由基础模型决定，"
            "模型之间的鲁棒性差异大于单一填补方法带来的平均差异。第二，不存在脱离基础模型和数据集的统一最优填补方法，"
            "SAITS和GP-RBF在总体上更稳定，Mean等强平滑方法在特定模型上仍可能取得较低预测误差。第三，缺失率、点级填补误差"
            "和历史结构扰动共同影响下游预测，其中结构扰动更适合作为连接填补质量与预测变化的中间量。"
        ),
    )

    doc.save(str(docx_path))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Rework manuscript Section 4.2 from the main experiment CSV.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Main experiment summary CSV.")
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX, help="Manuscript docx to update.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR, help="Figure output directory.")
    parser.add_argument("--data-dir", type=Path, default=DEFAULT_DATA_DIR, help="Summary CSV output directory.")
    parser.add_argument("--figures-only", action="store_true", help="Only generate summary CSVs and figures.")
    parser.add_argument("--docx-only", action="store_true", help="Only replace Section 4.2 in the manuscript.")
    return parser


def main() -> None:
    args = build_parser().parse_args()
    if args.docx_only and args.figures_only:
        raise SystemExit("--figures-only and --docx-only cannot be used together.")

    if not args.docx_only:
        result = generate_outputs(args.input, args.output_dir, args.data_dir)
        print(f"Imputed records: {len(result['records'])}")
        print(f"Figures: {args.output_dir}")
        print(f"Summary data: {args.data_dir}")
        print(f"Data manifest: {DATA_MANIFEST}")

    if not args.figures_only:
        update_docx(args.docx, args.output_dir, args.input)
        print(f"Updated docx: {args.docx}")


if __name__ == "__main__":
    main()
